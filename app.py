
from fastapi import FastAPI, Request, UploadFile, File, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import uvicorn, os, io, re, json
from collections import Counter
from typing import List, Dict
from docx import Document
import pandas as pd

APP_DIR = os.path.dirname(__file__)
UPLOAD_DIR = os.path.join(APP_DIR, "uploads")
EXPORT_DIR = os.path.join(APP_DIR, "exports")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

app = FastAPI(title="Citation Extractor")
app.mount("/static", StaticFiles(directory=os.path.join(APP_DIR, "static")), name="static")
templates = Jinja2Templates(directory=os.path.join(APP_DIR, "templates"))

# ---------------- Citation extraction utils ----------------

RE_YEAR = r"(?:19|20)\d{2}[a-z]?"
RE_NAME = r"[A-Z][A-Za-z'\-]+"

def _clean_piece(s: str) -> str:
    s = s.strip()
    # remove leading connector phrases commonly used in citations
    s = re.sub(r"^(?:see|e\.g\.|cf\.|e\.g|i\.e\.,?|for example|see also)\s*[:,]?\s*", "", s, flags=re.I)
    s = s.strip(" .;,")
    return s

def extract_apa_parenthetical(text: str) -> List[str]:
    results = []
    # Grab each (...) that contains a year
    for m in re.finditer(r"\(([^()]*?(?:19|20)\d{2}[^()]*)\)", text):
        inside = m.group(1)
        for part in re.split(r";", inside):
            piece = _clean_piece(part)
            if re.search(RE_YEAR, piece):
                # remove page refs ", p. 23" or "pp. 2-4"
                piece = re.sub(r",?\s*p+p?\.?\s*\d+(\-\d+)?", "", piece, flags=re.I)
                # trim content after year
                piece = re.sub(rf"({RE_YEAR}).*$", r"\1", piece)
                # normalize "and" to "&"
                piece = re.sub(r"\s+and\s+", " & ", piece)
                # collapse spaces
                piece = re.sub(r"\s{2,}", " ", piece).strip(" ,;")
                results.append(piece)
    return results

def extract_apa_narrative(text: str) -> List[str]:
    results = []
    # Pattern: Author ... (2019)
    for m in re.finditer(rf"((?:{RE_NAME}(?:\s*&\s*{RE_NAME})?|{RE_NAME}\s+et al\.))\s*\(\s*({RE_YEAR})\s*\)", text):
        name = m.group(1)
        year = m.group(2)
        name = re.sub(r"\s+and\s+", " & ", name)
        results.append(f"{name}, {year}")
    return results

def extract_ieee(text: str) -> List[str]:
    nums = []
    for m in re.finditer(r"\[(\d+(?:\s*[\-,–]\s*\d+|\s*,\s*\d+)*)\]", text):
        blob = m.group(1)
        parts = re.split(r"\s*,\s*", blob)
        for p in parts:
            if re.search(r"[\-–]", p):
                a,b = re.split(r"[\-–]", p)
                try:
                    a = int(a.strip()); b = int(b.strip())
                    if a<=b:
                        nums.extend([str(x) for x in range(a,b+1)])
                except: 
                    continue
            else:
                nums.append(p.strip())
    uniq = sorted(set(nums), key=lambda x: int(x) if x.isdigit() else x)
    return [f"[{n}]" for n in uniq]

def extract_citations_from_text(text: str) -> Dict[str, List[str]]:
    apa_p = extract_apa_parenthetical(text)
    apa_n = extract_apa_narrative(text)
    ieee  = extract_ieee(text)
    def dedupe(seq):
        seen=set(); out=[]
        for s in seq:
            s=s.strip()
            if s and s not in seen:
                out.append(s); seen.add(s)
        return out
    return {
        "APA_parenthetical": dedupe(apa_p),
        "APA_narrative": dedupe(apa_n),
        "IEEE_numeric": dedupe(ieee),
    }

def read_docx_text(file_bytes: bytes) -> str:
    bio = io.BytesIO(file_bytes)
    doc = Document(bio)
    texts=[]
    for p in doc.paragraphs:
        texts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)

# ---------------- Routes ----------------

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/extract", response_class=HTMLResponse)
async def extract(request: Request, file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Please upload a .docx file")
    content = await file.read()
    text = read_docx_text(content)
    buckets = extract_citations_from_text(text)

    rows = []
    from collections import Counter
    for kind, items in buckets.items():
        for it in items:
            rows.append({"Type": kind.replace("_"," "), "Citation": it})
    counts = Counter()
    for r in rows:
        counts[r["Citation"]] += 1
    for r in rows:
        r["Count"] = counts[r["Citation"]]

    import pandas as pd, os
    df = pd.DataFrame(rows).drop_duplicates().sort_values(["Type","Citation"]).reset_index(drop=True)

    export_name = f"citations_{os.path.splitext(file.filename)[0]}.csv"
    export_path = os.path.join(EXPORT_DIR, export_name)
    df.to_csv(export_path, index=False)

    return templates.TemplateResponse("results.html", {
        "request": request,
        "filename": file.filename,
        "buckets": buckets,
        "rows": df.to_dict(orient="records"),
        "download_name": export_name
    })

@app.get("/download/{name}")
async def download(name: str):
    path = os.path.join(EXPORT_DIR, name)
    if not os.path.exists(path):
        raise HTTPException(404, "File not found")
    return FileResponse(path, filename=name, media_type="text/csv")

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8001, reload=True)
